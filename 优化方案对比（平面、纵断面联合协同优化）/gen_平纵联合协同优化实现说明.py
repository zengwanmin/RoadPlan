# -*- coding: utf-8 -*-
"""生成《平纵联合协同优化实现说明》.docx
逐段说明: 两套方案共用同一成本/能耗模型/约束/算法(IJS)的前提下,
"平纵联合协同"是如何在代码与公式层面得出并体现的。全部内容取自本文件夹
run_joint.py / objective_joint.py / objective.py / params.py / algorithms.py /
data_loader.py / safety.py, 未杜撰。"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

OUT = "平纵联合协同优化实现说明.docx"
doc = Document()

normal = doc.styles["Normal"]
normal.font.name = "宋体"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = RGBColor(0, 0, 0)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    return p


def para(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def bullet(text):
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)
    return p


def code(text):
    """等宽代码块。"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r._element.rPr.rFonts.set(qn('w:eastAsia'), "Consolas")
    p.paragraph_format.left_indent = Pt(12)
    return p


def formula(text):
    """公式行(居中、楷体感)。"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(11)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for j, htext in enumerate(headers):
        c = t.rows[0].cells[j]
        c.text = ""
        run = c.paragraphs[0].add_run(htext)
        run.bold = True
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = str(val)
    return t


# ============================ 正文 ============================
doc.add_heading("平纵联合协同优化实现说明", level=0)
para("——在同一成本/能耗模型、同一约束、同一算法(IJS)下, “平纵联合协同”是如何实现与体现的",
     bold=True)
para("说明: 本文所有代码片段与公式均取自本文件夹内的 run_joint.py、objective_joint.py、"
     "objective.py、params.py、algorithms.py、data_loader.py、safety.py, 未作杜撰。凡论文未给"
     "具体取值、由工程标定者, 均按 params.py 中的原始注释如实标注。")

# ---- 0 ----
h("0  一句话概览", level=1)
para("两套方案(两阶段 / 平纵联合)共用同一套 C(全周期成本)、E(全周期能耗)公式、同一套"
     "几何约束、同一个 IJS 求解器。它们唯一的区别在于“决策变量怎么组织、怎么搜”:")
bullet("两阶段: 先只搜平面(25维)→冻结→再只搜纵断面(225维), 分两次串联寻优。")
bullet("平纵联合: 把平面25维与纵断面225维拼成一个250维向量, 在同一次 IJS 寻优里一起搜, "
       "平面走向与纵断面坡度相互影响、共同决定同一个标量目标 F。")
para("因此“协同”不是换了一套模型或公式, 而是体现在“决策变量的耦合”与“单目标标量化的"
     "联合评价”这两处代码机制上。下面从数据到目标逐层说明。")

# ---- 1 ----
h("1  两套方案共用的底座(模型/约束/算法完全一致)", level=1)
para("先明确“同一”指的是哪些东西。以下模块两个文件夹逐字节相同(或仅调用方式不同), "
     "它们保证了两套方案的可比性:")
table(
    ["模块", "内容", "关键公式(林坤锐学位论文式号)"],
    [
        ["data_loader.py", "实测轨迹→里程/地面高程/平面坐标", "式3.1-3.4 (大圆距离+平面直角坐标)"],
        ["objective.py", "土方C_TU、平面全周期C_PING、油耗E_fuel、电耗E_ele、熵权", "式4.3 / 式3.41-3.55 / 式4.5-4.18 / 式5.3-5.4"],
        ["safety.py", "边坡失稳危险度Q", "式6.1, 表6.7"],
        ["algorithms.py", "IJS(Tent+Levy+DE)统一求解器", "式47-59"],
        ["params.py", "全部公式常量/技术标准/标定值", "表3.2/4.4/6.4/6.6 等"],
    ],
)
para("成本与能耗的顶层定义(objective.py / objective_joint.py 一致):")
formula("C = C_PING + C_TU = (CR + CB + CS + CQ) + C_TU     (式3.41 + 式4.3)")
formula("E = E_fuel + E_ele                                (式4.4 / 4.26, 全周期货币化)")
para("约束(两套方案同一, 见 objective_joint.py 惩罚项): 平曲线半径 R≥400m(表3.2极限值)、"
     "纵坡 |i|≤4%(式4.27, 表4.4)、相邻纵坡代数差≤3%(式4.28-4.29竖曲线近似)。求解器同为"
     "VARIANTS[“V5_IJS”](Tent混沌初始化+Levy飞行+差分进化, algorithms.py)。")

# ---- 2 ----
h("2  “协同”的实现要点一: 决策变量的一体化编码", level=1)
para("这是平纵联合区别于两阶段的第一处、也是最本质的地方。见 objective_joint.py 顶部定义:")
code("CORRIDOR_HALF_W = 800.0   # 走廊带半宽(m)\n"
     "N_CTRL = 25               # 平面控制点个数\n"
     "M_PROF = 225              # 纵断面变坡点个数\n"
     "# 决策向量 x ∈ [0,1]^(N_CTRL + M_PROF) = [0,1]^250")
para("一个归一化到[0,1]的250维向量 x 同时编码了平面与纵断面:")
formula("x[0:25]   → 平面25个控制点的法向偏移 δ_k ∈ [-W, +W],  W=800m")
formula("x[25:250] → 沿【新平面线位】等分的225个变坡点高程调整")
para("解码函数 decode_joint(x, pc) 的关键在于: 平面分量先解出, 纵断面分量建立在平面结果"
     "之上, 二者在一次解码中前后相连——这正是“协同”在代码里的落点:")
code("delta = (x[:25] - 0.5) * 2 * W        # 平面法向偏移 δ_k\n"
     "xx, yy = build_plane(pc, delta)       # 三次样条→新平面线形(平面走向由此变)\n"
     "L_new, R = _plane_metrics(xx, yy)     # 新里程 L_new、新曲率半径序列 R\n"
     "# 沿【新线位】等弧长布桩 → 新桩号 sta\n"
     "sta = 等分(新线位弧长, 225)\n"
     "# 新桩号点地面高程 = 最近实测中线点高程(KDTree最近邻, 数据受限下的近似)\n"
     "gz_new = pc[\"gz_meas\"][最近邻索引(sta点)]\n"
     "design_z = gz_new + (x[25:] - 0.5) * 2 * amp   # 纵断面设计高程")
para("关键点在于: 纵断面所依附的桩号 sta、地面高程 gz_new、乃至里程 L_new, 全部由平面分量"
     "x[0:25]决定。平面一动, 纵断面的“地基”就跟着变。因此在联合向量里, 平面与纵断面不是"
     "两个独立子问题, 而是通过 decode_joint 强耦合在一起。这与两阶段“平面先定死、纵断面"
     "在固定线位上单独优化”形成本质区别。")

# ---- 2.1 平面25维 ----
h("2.1  平面为什么是25维: 25个平面控制点的“横向摆动量”", level=2)
para("平面25维 = 沿路线等弧长布置的 25 个平面控制点, 每个控制点贡献 1 个决策变量, 即该点"
     "垂直于路线方向(法向)的横向偏移量 δ_k。含义与生成方式如下(make_plane_context / "
     "build_plane / decode_joint):")
bullet("布点: 对实测中线按弧长等分取 N_CTRL=25 个控制点。本段全长约22.46km, 故相邻控制点"
       "间距 ≈ 22462.5 / (25-1) ≈ 936m。即大约每0.94km设一个可摆动的平面控制点。")
bullet("每维的物理量: 第k个控制点的法向偏移 δ_k = (x[k]-0.5)·2·W, 走廊带半宽 W=800m, "
       "故 δ_k ∈ [-800m, +800m]。x[k]=0.5 表示不偏(落在实测中线上), <0.5 向一侧摆、>0.5 "
       "向另一侧摆。")
bullet("从25个点到整条平面线: 把25个偏移后的控制点用三次样条(splprep, k=3)插值加密成600点"
       "的光滑曲线, 作为新平面线形; 首末两点固定(端点不动, 保证接线)。所以真正的自由度是"
       "中间那些控制点, 25这个数字控制“平面能有多少处独立摆动”。")
bullet("为什么取25而非更多: 25点≈每0.94km一个自由度, 既能表达路线走向的整体优化(取直、"
       "顺势绕避), 又不至于让平面产生高频抖动(样条会被过多控制点带偏), 同时与 GapB 设置"
       "保持一致(见 objective_joint.py 注释)。")
para("一句话: 平面25维 = “允许路线在25个断面处横向摆动, 每处摆动范围±800m”, 由此决定新的"
     "平面走向, 进而决定里程 L_new 与曲率半径 R。")

# ---- 2.2 纵断面225维 ----
h("2.2  纵断面为什么是225维: 225个变坡点的“竖向抬降量”", level=2)
para("纵断面225维 = 沿【新平面线位】等弧长布置的 225 个变坡(桩号)点, 每个点贡献 1 个决策"
     "变量, 即该桩号处设计高程相对地面的抬升/下降量。含义与生成方式如下(decode_joint):")
bullet("布桩: 在新平面线位上按弧长等分 M_PROF=225 个桩号。新线位长约22.34km(现状δ=0时), "
       "故桩号间距 ≈ 22342.8 / (225-1) ≈ 99.7m ≈ 100m。即约每100m一个变坡点, 与论文纵断面"
       "按桩号离散、连续坡段近似为直线坡的处理(§4.2 假设)一致。")
bullet("每维的物理量: 第i个桩号的设计高程 design_z_i = gz_new_i + (x[25+i]-0.5)·2·amp, 其中"
       "gz_new_i 是该桩号处的地面高程, amp=max(地面起伏范围×0.6, 10)m 为允许调整幅度。"
       "x=0.5 表示贴地(挖填为0), 偏离0.5即抬升(填方)或下降(挖方)。")
bullet("为什么是225: 桩号间距≈100m是纵断面设计的常用离散尺度; 22.46km÷100m≈225, 于是取"
       "225个变坡点。相邻两点高程之差决定该坡段纵坡 g_i=ΔH/ΔP(式4.19), 纵坡再决定土方、"
       "能耗(坡度阻力)与竖曲线约束。")
para("一句话: 纵断面225维 = “允许纵断面在约每100m一个桩号处独立抬降, 幅度±amp”, 由此决定"
     "填挖高度、纵坡与竖曲线, 进而决定土方费 C_TU、养护费 CQ 与车流能耗 E。")

# ---- 2.3 维度小结表 ----
h("2.3  两组维度一览", level=2)
table(
    ["", "平面分量", "纵断面分量"],
    [
        ["维数", "25 (N_CTRL)", "225 (M_PROF)"],
        ["决策向量位置", "x[0:25]", "x[25:250]"],
        ["每维物理含义", "控制点法向横向偏移 δ_k", "变坡点设计高程抬降量"],
        ["取值范围", "δ_k ∈ [-800m, +800m] (走廊带±W)", "±amp = ±max(起伏×0.6,10) m"],
        ["布点方式", "沿实测中线等弧长25点", "沿新平面线位等弧长225点"],
        ["相邻间距(本段)", "≈ 936 m/点", "≈ 100 m/点"],
        ["x=0.5 的含义", "不偏移(落在实测中线上)", "贴合地面(挖填为0)"],
        ["主要影响的指标", "里程 L、曲率半径 R → CR/CB/CS", "纵坡 i、填挖 h → C_TU/CQ/E/Q"],
    ],
)
para("合计 25 + 225 = 250 维, 即联合决策向量 x ∈ [0,1]^250。平面25维“调走向”、纵断面225维"
     "“调竖向”, 二者在同一向量中被 IJS 一次性协同优化。")

# ---- 3 ----
h("3  “协同”的实现要点二: 用【新里程/新桩号】联合计算 C 与 E", level=1)
para("解码得到(新平面, 新桩号 sta, 新地面高程 gz_new, 设计高程 design_z)后, objectives_joint "
     "用这些【联合产生的量】统一算成本与能耗。这一步让平面的贡献(里程)与纵断面的贡献"
     "(填挖/坡度)进入同一个目标, 是“协同评价”的核心。")

h("3.1  目标一 成本 C(元): 平面决定里程, 里程与纵断面共同决定各项费用", level=2)
para("C = C_PING + C_TU。其中(objective.py 的 lcc_ping / earthwork_cost, 传入的是新里程 L_new "
     "与新桩号):")
bullet("占地费 CR(式3.42-3.44): CR = 占地面积×单价, 占地面积 = L_new × 路基宽/666.67(亩)。"
       "→ 直接随平面里程 L_new 变化。")
bullet("桥隧费 CB(式3.45-3.51, 按里程综合造价): CB = β · L_new(km), β=4.2948e7 元/km(由现状"
       "反标定, 见 params.py)。CB 约占 C 的94%, 故里程一变 C 就显著变——平面的主导作用在此。")
bullet("基建费 CS(式3.52-3.54): CS = 路基(200元/m×L) + 路面(3e4元/km×L), 亦随里程。")
bullet("养护费 CQ(式3.55): 基础养护 + 交通量项 + 边坡项(Σ填挖高×坡率×50), 30年折现; "
       "含纵断面填挖高度 h 的贡献。")
bullet("土方费 C_TU(式4.3): C_TU = 调运费 + Ks·挖方量 + Kh·填方量, 挖填量由 |design_z - gz_new| "
       "沿新桩号积分而来 → 纵断面分量 x[25:] 的直接贡献。")
para("可见 C 里既有“平面里程项”(CR/CB/CS)又有“纵断面填挖项”(C_TU/CQ), 两类变量在同一个 C "
     "中相加, 一次评价即完成协同。")

h("3.2  目标二 能耗 E(元): 沿新线位的油电混合车流全周期能耗", level=2)
para("E = E_fuel + E_ele, 单车单程能耗按逐坡段力学模型积分, 再乘车流与全周期折现系数:")
formula("F_ext = F_air + F_r + F_g   (式4.9-4.11: 空气/滚动/坡度阻力)")
formula("HP_ex = F_ext·v / 736,  UFC = φ·(HP_in + HP_ex)   (式4.6-4.7 油耗率)")
formula("E_fuel = AADT·n1·(ml/1000)·K_f·lc_factor·CALIB    (油车全周期)")
formula("E_ele  = AADT·n2·kwh·Z_Q·lc_factor·CALIB           (电车全周期, 下坡再生回收 式4.17)")
para("其中坡度阻力 F_g = m·g·sinθ, θ=arctan(纵坡), 纵坡由 design_z 与 sta 求得。里程越短、"
     "纵坡越平顺, 单程能耗越低。由于 sta 与 design_z 都来自联合解码, 平面(缩里程)与纵断面"
     "(平顺坡度)对 E 的改善也是在同一次评价里叠加的。lc_factor = 365×30年等额年金现值系数"
     "(5%折现), CALIB=0.6792 为对齐论文表6.8 现状 E=8.95亿元的统一标定系数(params.py)。")

# ---- 4 ----
h("4  “协同”的实现要点三: 统一的约束惩罚(平面+纵断面同时受约束)", level=1)
para("联合模型里, 平面与纵断面的几何约束被写进同一个 penalty, 与两目标一起决定 F。这保证"
     "协同优化在压里程/降能耗时不会违反规范(objective_joint.py):")
code("pen = 0\n"
     "# (a) 平面: 最小平曲线半径 R ≥ 400m (表3.2)\n"
     "pen += Σ max((400 - R)/400, 0) · 5e7\n"
     "# (b) 纵断面: 纵坡 |i| ≤ 4% (式4.27)\n"
     "pen += Σ max(|i| - 0.04, 0) · 1e9\n"
     "# (c) 竖曲线: 相邻纵坡代数差 ≤ 3% (式4.28-4.29)\n"
     "pen += Σ max(|Δi| - 0.03, 0) · 5e8")
para("(a)约束的是平面分量 x[0:25]生成的曲率半径 R, (b)(c)约束的是纵断面分量 x[25:]生成的"
     "纵坡。三者累加成一个 penalty——平面和纵断面在“合规性”上也被绑在一起评价。惩罚系数"
     "(5e7/1e9/5e8)为足够大的工程标定量级, 由 run.log 中 pen→0 验证最终方案落在可行域内。")

# ---- 5 ----
h("5  “协同”的实现要点四: 熵权法标量化 → 单目标 IJS 寻优", level=1)
para("双目标(C, E)要交给单目标的 IJS, 需先标量化。联合模型用熵权法客观定权后合成一个标量 F:")
formula("F = wC·(C/C_ref) + wE·(E/E_ref) + penalty/C_ref     (式5.3-5.4)")
para("三点说明:")
bullet("归一化: C/C_ref、E/E_ref 把两个量纲不同、量级不同的目标压到同一尺度(~1量级), "
       "penalty/C_ref 也归一后可加。C_ref、E_ref 取初始种群 C、E 的均值。")
bullet("权重客观化: wC、wE 由初始种群的(C,E)分布经信息熵与差异系数算出(entropy_weights, "
       "式5.3-5.4), 非人为指定。本次联合实验算得 wC=0.674、wE=0.326(见 run.log)。")
bullet("协同的落点: F 把“平面缩里程带来的 C/E 改善”和“纵断面平顺带来的 C/E 改善”合成到"
       "同一个标量里。IJS 每次评价一个250维个体, 得到的就是平纵共同作用下的 F——算法据此"
       "同时调整平面与纵断面, 实现真正的一体化协同。")

# ---- 6 ----
h("6  IJS 求解器如何在250维上一起搜平面与纵断面", level=1)
para("IJS(algorithms.py 的 run + VARIANTS[“V5_IJS”])是无约束单目标优化器, 对250维向量整体"
     "更新, 不区分哪几维是平面、哪几维是纵断面——这正是“一起搜”的机制。三大改进组件:")
bullet("Tent混沌初始化(式47): 在随机种群上做混沌扰动择优, 提升初始多样性。")
bullet("洋流/主动/被动运动(式48-53): 由时间控制函数 Ar=(1-t/Max)(2rand-1)在探索与开发间切换。")
bullet("Levy飞行(式54-55)+差分进化DE(式56-58): 大跳变+差分变异增强跳出局部最优能力。")
para("每次迭代对整条250维向量施加上述算子, 再用标量 F 做贪婪选择。也就是说, 平面25维与"
     "纵断面225维在同一次种群更新中被同步扰动、同步评价、同步保留, 没有先后之分。这就是"
     "“平纵一体化协同寻优”在算法层面的含义。")

# ---- 7 ----
h("7  端到端流程(run_joint.py 主程序)", level=1)
para("把上述环节串起来, 联合实验的完整执行顺序如下(run_joint.py main):")
bullet("① 载入实测轨迹(data_loader.load_alignment)→ make_plane_context 预计算平面控制点、"
       "法向、KDTree、原始里程。")
bullet("② 构造初始种群 base(200×250): 平面分量给足±全走廊带初始幅度(注释原文: “避免平面"
       "子空间(仅N_CTRL维)在高维联合搜索中被纵断面(M_PROF维)淹没”), 纵断面分量随机。")
bullet("③ 用初始种群的(C,E)算熵权 wC/wE 与参考尺度 C_ref/E_ref(entropy_weights)。")
bullet("④ 评价三方案(均在同一联合模型下):")
para("     M-A 现状方案: 平面δ=0(实测中线) + 纵断面按0.5km尺度平滑地面线(人工粗放, 未优化), "
     "作为基线。", )
para("     M-B 单目标成本最优: 平纵联合优化, 仅 min C(wC=1, wE=0)。", )
para("     M-C 平纵联合双目标(本文): 平纵联合优化, min C 与 min E 协同 + 熵权决策。", )
para("     M-B→M-C 的差 = “引入能耗协同优化”的净贡献。")
bullet("⑤ 权重扫描 w_C=0.1~0.9 生成 Pareto 参考前沿(图C1), 熵权点即 M-C。")
bullet("⑥ 输出 joint_results.json / 表C1 / 表C2 / 图。")
para("本次联合实验结果(run.log): M-A C=10.2400亿/E=8.9497亿 → M-C C=9.7901亿/E=8.5487亿, "
     "里程 22.343→21.453km(缩短3.98%), Rmin=968m(≥400合规), pen=0。")

# ---- 8 ----
h("8  与两阶段的实现差异(一表看清“协同”体现在哪)", level=1)
table(
    ["环节", "两阶段(先平面后纵断面)", "平纵联合协同"],
    [
        ["决策变量", "分两批: 先25维平面, 再225维纵断面", "一个250维向量(平面25+纵断面225)"],
        ["寻优次数", "两次串联(Stage1→冻结→Stage2)", "一次(平纵同时搜)"],
        ["平面与纵断面关系", "平面先定死, 纵断面在固定线位上优化", "平面动→纵断面地基随动, 强耦合"],
        ["标量目标 F", "Stage1: 平面LCC; Stage2: wC·C+wE·E+pen", "wC·C+wE·E+pen (平纵合成同一F)"],
        ["熵权 wC/wE", "0.417 / 0.583 (仅对纵断面采样定权)", "0.674 / 0.326 (对联合向量采样定权)"],
        ["成本/能耗模型", "同一套(objective.py)", "同一套(objective.py)"],
        ["约束/惩罚", "同一套(R≥400, |i|≤4%, |Δi|≤3%)", "同一套"],
        ["求解器", "同一 IJS(V5)", "同一 IJS(V5)"],
    ],
)
para("结论: 模型、约束、算法三者完全相同; “协同”仅体现在“决策变量一体化编码 + 单目标"
     "标量联合评价 + IJS对250维整体寻优”这一条链路上。")

# ---- 9 ----
h("9  数据受限下的诚实声明(不杜撰)", level=1)
bullet("面状地形缺失: 数据.xlsx 仅有实测中线一条轨迹的地面高程, 无面状DEM。平面横向偏移后"
       "新点位地面高程用“最近实测中线点高程”近似(cKDTree最近邻)。此近似在走廊带内、地形"
       "沿纵向变化为主时成立, 属数据限制下的必要处理, 已在 objective_joint.py 模块头声明。")
bullet("桥隧/能耗标定: 论文桥隧费用系数(a1~a4等)与能源单价均标注“由工程定”未给取值; 本文"
       "以单位里程综合造价 β=4.2948e7元/km、能耗标定 CALIB=0.6792, 由现状方案反标定使 C、E "
       "命中论文表6.8(10.24亿/8.95亿)。系数对所有方案同一, 不改变随线形变化的相对趋势。")
bullet("惩罚系数为工程标定量级, 靠 run.log 的 pen→0 验证可行, 而非精确物理值。")

# ---- 10 ----
h("10  一句话总结", level=1)
para("平纵联合协同 = 把平面(25维走向)与纵断面(225维高程)拼成一个250维决策向量, 经 "
     "decode_joint 强耦合解码(平面定里程与桩号、纵断面在其上布高程), 用同一套 C/E 公式与"
     "约束算出单一标量 F=wC·C/C_ref+wE·E/E_ref+pen/C_ref, 交给同一个 IJS 对整条向量一次性"
     "寻优。模型、约束、算法都没变, 变的只是“变量组织方式”——这就是协同的全部实现与体现。",
     bold=True)

doc.save(OUT)
print("full doc saved ->", OUT)
