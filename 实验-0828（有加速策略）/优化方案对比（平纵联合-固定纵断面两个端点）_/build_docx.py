# -*- coding: utf-8 -*-
"""把 md 转 docx，并将 ==...== 标记的文字高亮为绿色。
做法: 先把 ==...== 换成哨兵 -> pandoc 转 docx -> python-docx 遍历 run 上色。"""
import re, subprocess, sys
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

MD = "优化方案对比实验说明与分析总结.md"
DOCX = "优化方案对比实验说明与分析总结.docx"
OPEN, CLOSE = "〖HL〗", "〖/HL〗"  # 〖HL〗 〖/HL〗

# 1. 读 md, 把 ==x== 换成哨兵
src = open(MD, encoding="utf-8").read()
tmp = re.sub(r"==(.+?)==", OPEN + r"\1" + CLOSE, src, flags=re.S)
open(".tmp_hl.md", "w", encoding="utf-8").write(tmp)

# 2. pandoc 转 docx
subprocess.run(["pandoc", ".tmp_hl.md", "-o", DOCX], check=True)

# 3. python-docx: 对含哨兵的段落逐段处理, 高亮哨兵之间的文字
doc = Document(DOCX)
for para in doc.paragraphs:
    full = "".join(r.text for r in para.runs)
    if OPEN not in full:
        continue
    # 重建该段落: 清空 runs, 按标记状态重写
    for r in list(para.runs):
        r.text = ""
    # 解析 full 为片段序列
    parts = re.split(r"(" + re.escape(OPEN) + r"|" + re.escape(CLOSE) + r")", full)
    base = para.runs[0] if para.runs else para.add_run()
    base.text = ""
    hl = False
    first = True
    for seg in parts:
        if seg == OPEN:
            hl = True; continue
        if seg == CLOSE:
            hl = False; continue
        if seg == "":
            continue
        run = base if first else para.add_run()
        first = False
        run.text = seg
        if hl:
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
doc.save(DOCX)

# 4. 校验
import zipfile
x = zipfile.ZipFile(DOCX).read("word/document.xml").decode("utf-8")
print("green highlight runs:", x.count('w:val="green"'))
print("tables:", x.count("<w:tbl>"))
print("residual sentinel:", (OPEN in x) or (CLOSE in x))
